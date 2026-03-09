#!/usr/bin/env python3
"""Generate Wynisco AI Matching Service — Technical Document (.docx)"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime
import diagram_generator

doc = Document()

# ── Styles ──────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f"Heading {level}"]
    h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers), style="Light Grid Accent 1")
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(9)
    for row_data in rows:
        row = t.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                if w:
                    row.cells[i].width = Cm(w)
    return t


def add_code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(f" {text}")
    else:
        p.add_run(text)


def add_diagram(path, width=Inches(6.0)):
    doc.add_picture(path, width=width)


def add_note(text):
    """Add an italic note/subtext paragraph in gray."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x77)


# Pre-generate all diagram images
diagrams = diagram_generator.generate_all()


# ══════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════
spacer = doc.add_paragraph()
spacer.paragraph_format.space_before = Pt(180)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("AI Resume Matching Service")
run.font.size = Pt(32)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Integration, Database Schema & Deployment Guide")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x77)

spacer2 = doc.add_paragraph()
spacer2.paragraph_format.space_before = Pt(40)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(
    f"Version 1.0 | {datetime.date.today().strftime('%B %d, %Y')}"
).font.size = Pt(12)

meta2 = doc.add_paragraph()
meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta2.add_run("Confidential — Internal Use Only").font.color.rgb = RGBColor(
    0xCC, 0x00, 0x00
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════
doc.add_heading("Table of Contents", level=1)
toc_items = [
    (0, "1. What This Service Does"),
    (0, "2. How It Fits Into the Existing Architecture"),
    (0, "3. Database Schema & Relationships"),
    (1, "3.1 Existing Tables the Service Reads"),
    (1, "3.2 New Tables the Service Creates"),
    (1, "3.3 How New Tables Connect to Existing Tables"),
    (1, "3.4 Complete Data Chain (Worked Example)"),
    (1, "3.5 Database Ownership & Migration Safety"),
    (0, "4. AI Matching Pipeline"),
    (1, "4.1 Three-Layer Pipeline"),
    (1, "4.2 Two-Layer Caching"),
    (0, "5. Deployment"),
    (1, "5.1 Environment Variables"),
    (1, "5.2 Local Development"),
    (1, "5.3 Docker & Cloud Run"),
    (0, "6. Cost Estimates"),
]
for indent_level, item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.runs[0].font.size = Pt(10)
    if indent_level > 0:
        p.paragraph_format.left_indent = Cm(1)
    p.runs[0].bold = indent_level == 0

doc.add_page_break()


# ══════════════════════════════════════════════════════════
# 1. WHAT THIS SERVICE DOES
# ══════════════════════════════════════════════════════════
doc.add_heading("1. What This Service Does", level=1)

doc.add_paragraph(
    "Today, when a consultant at Wynisco needs to find the best jobs for a candidate, "
    "they manually review job postings, compare requirements against the candidate's "
    "resume, and make subjective judgments about fit. This process is slow (hours per "
    "candidate), inconsistent (different consultants evaluate differently), and doesn't "
    "scale as the number of candidates and jobs grows."
)

doc.add_paragraph(
    "The AI Resume Matching Service automates this process. It is a new microservice "
    "that reads candidate resumes and job descriptions from the existing Wynisco database, "
    "uses AI to analyze how well each candidate fits each job, and produces ranked match "
    "results with plain-English explanations of why a match is good or bad."
)

doc.add_paragraph(
    "Critically, this service does not replace any existing system. It runs alongside "
    "the existing SMO backend as a separate process. It connects to the same PostgreSQL "
    "database that SMO already uses, reads existing candidate and job data, and writes "
    "only to 3 new tables that it owns. The SMO backend, SMO frontend, and all existing "
    "workflows continue to operate exactly as before."
)

doc.add_heading("Service at a Glance", level=2)

add_table(
    ["Aspect", "Detail"],
    [
        ["Service Name", "auto-matching-service"],
        ["Technology", "Python 3.11, FastAPI (async), SQLAlchemy 2.0"],
        ["Ports", "8002 (backend API), 3001 (dashboard UI)"],
        ["AI Models",
         "OpenAI text-embedding-3-small for converting text into numerical "
         "representations; Cerebras gpt-oss-120b for intelligent scoring and analysis"],
        ["Database",
         "Same PostgreSQL instance as SMO (smo_local). Reads 6 existing tables, "
         "writes to 3 new tables it owns."],
        ["Caching",
         "Two-layer: Redis for individual AI results (72hr TTL, optional) and "
         "database for full match results (24hr TTL). Works without Redis."],
        ["Dashboard",
         "Next.js web application for browsing match results, viewing AI "
         "explanations (WhyCards), tuning score weights, and listing candidates/jobs."],
    ],
    col_widths=[2.5, 13],
)

add_note(
    "The matching service is fully independent — it can be started, stopped, redeployed, "
    "or scaled without affecting the SMO backend or frontend in any way."
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════
# 2. HOW IT FITS INTO EXISTING ARCHITECTURE
# ══════════════════════════════════════════════════════════
doc.add_heading("2. How It Fits Into the Existing Architecture", level=1)

doc.add_paragraph(
    "The Wynisco platform already has two main components: the SMO Backend (FastAPI, "
    "port 8000) which handles all candidate, job, and application management, and the "
    "SMO Frontend (React, port 3000) which is the main user interface for consultants "
    "and admins."
)

doc.add_paragraph(
    "The matching service adds two new components to this picture: a Matching Backend "
    "(FastAPI, port 8002) that runs the AI pipeline, and a Matching Dashboard (Next.js, "
    "port 3001) that provides a dedicated UI for viewing and working with match results. "
    "Both new components plug into the existing infrastructure by sharing the same "
    "PostgreSQL database — they do not require any changes to the existing SMO code."
)

add_diagram(diagrams["architecture"])

doc.add_heading("How Each Connection Works", level=2)

doc.add_paragraph(
    "The diagram above shows all the services and how they communicate. Here is a "
    "detailed breakdown of each connection — what data moves, in which direction, "
    "and why it matters:"
)

add_table(
    ["Connection", "Direction", "What Happens"],
    [
        ["Matching Backend → PostgreSQL",
         "Reads existing data",
         "The matching service connects to the same PostgreSQL database that SMO uses. "
         "It reads candidate profiles (resume text, skills, experience, salary expectations), "
         "job postings (descriptions, requirements, locations), and employer information "
         "(company names, blacklist flags). This is a read-only relationship — the matching "
         "service never modifies any SMO-owned data."],
        ["Matching Backend → PostgreSQL",
         "Writes new data",
         "The service writes to 3 tables it owns: candidate_embeddings (AI fingerprints of "
         "resumes), job_embeddings (AI fingerprints of job descriptions), and match_scores_v2 "
         "(scored results with explanations). These tables live in the same database but are "
         "completely separate from SMO tables."],
        ["Matching Backend → OpenAI API",
         "Outbound HTTPS",
         "When the service needs to create or update an embedding, it sends the resume or job "
         "description text to OpenAI's embedding API and receives back a vector of 1,536 numbers "
         "that represents the meaning of that text. This is a one-time operation per document — "
         "the result is stored and reused until the source text changes."],
        ["Matching Backend → Cerebras API",
         "Outbound HTTPS",
         "For each candidate-job pair that passes the earlier filtering stages, the service sends "
         "both the candidate's profile and the job details to Cerebras' LLM. The AI reads both "
         "sides and returns a structured analysis: a score (0-100), strengths, skill gaps, "
         "experience fit, salary fit, and a recommendation. This is packaged as a 'WhyCard'."],
        ["Matching Backend → Redis",
         "Cache read/write",
         "Before calling Cerebras for a candidate-job pair, the service checks Redis for a "
         "cached result. If found, it skips the API call entirely (saving ~2-3 seconds and "
         "API cost). Redis is optional — if not configured, the system still works using "
         "database-only caching."],
        ["Dashboard → Backend",
         "REST API",
         "The Next.js dashboard communicates with the matching backend via REST API calls. "
         "A built-in proxy rewrites /api/* requests to the backend (port 8002), so the "
         "frontend doesn't need to know the backend's address directly."],
        ["SMO Backend ↔ Matching",
         "No direct connection",
         "The two backends never call each other's APIs. They are completely decoupled services "
         "that happen to share a database. This means either service can be updated, restarted, "
         "or redeployed independently without affecting the other."],
    ],
    col_widths=[3, 2.5, 10],
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════
# 3. DATABASE SCHEMA & RELATIONSHIPS
# ══════════════════════════════════════════════════════════
doc.add_heading("3. Database Schema & Relationships", level=1)

doc.add_paragraph(
    "This section explains how the matching service's data relates to the existing "
    "Wynisco database. This is the most important part of this document because it "
    "answers the key question: how do the new tables connect to what already exists?"
)

doc.add_paragraph(
    "The short answer: the matching service connects to the same PostgreSQL database "
    "(smo_local) that the SMO backend already uses. It reads 6 existing tables to get "
    "candidate and job data, and it creates 3 new tables to store its AI results. The "
    "new tables reference existing data through candidate_id (linking to users) and "
    "job_id (linking to jobs), but they don't use hard database constraints — keeping "
    "the two services cleanly separated."
)


# ── 3.1 EXISTING TABLES ──
doc.add_heading("3.1 Existing Tables the Service Reads (Read-Only)", level=2)

doc.add_paragraph(
    "The matching service needs access to candidate and job data that already exists in "
    "the SMO database. Rather than duplicating this data, it reads directly from SMO's "
    "tables using read-only database mappings. This means the matching service always "
    "sees the latest data — when a consultant updates a candidate's skills or a new job "
    "is scraped, the matching service picks up the changes automatically."
)

doc.add_paragraph(
    "Here are the 6 tables it reads, what specific columns it uses from each, and why "
    "it needs them:"
)

add_table(
    ["Existing Table", "What the Matching Service Reads", "Why It Needs It"],
    [
        ["users",
         "id (UUID), email, first_name, last_name, role, pod_id",
         "Identifies which users are candidates (filters by role). The id is the primary "
         "key that links everything together — embeddings and match scores both reference "
         "this id. pod_id groups candidates into their cohort for filtering."],
        ["candidate_profiles",
         "user_id, resume_plain_text, skills[], job_titles[], "
         "years_of_experience, location, expected_salary_min/max, open_to_relocate",
         "This is the core input for matching. resume_plain_text is sent to OpenAI for "
         "embedding generation. skills, experience, location, and salary are used in hard "
         "filters (eliminate impossible matches) and sent to the LLM for detailed analysis."],
        ["jobs",
         "id, employer_id, job_title, job_description, location, "
         "state, city, salary, experience, created_at",
         "The other side of the match. job_description is sent to OpenAI for embedding. "
         "Title, location, salary, and experience requirements are used in filtering and "
         "LLM analysis. created_at filters out old postings (default: last 30 days)."],
        ["employers",
         "id, name, do_not_apply, h1b_filed, everified, location",
         "Every job links to an employer. The do_not_apply flag is critical — if an "
         "employer is blacklisted, all their jobs are excluded from matching. Company name "
         "is included in match results for display."],
        ["pods",
         "id (UUID), pod_name",
         "Pods are cohorts/batches of candidates. The matching dashboard lets you filter "
         "candidates by pod, so it needs the pod name for display."],
        ["universities",
         "id, school_name, city, state",
         "Some candidates have a university linked in their profile. This is used for "
         "display in match results — it helps consultants understand a candidate's background."],
    ],
    col_widths=[2.5, 5, 8],
)

add_note(
    "The matching service uses SQLAlchemy read-only mapped classes (defined in "
    "app/models/readonly.py) with extend_existing=True. This means it maps to the same "
    "physical tables as SMO but in a read-only mode — it cannot accidentally modify SMO data "
    "even if there's a code bug."
)


# ── 3.2 NEW TABLES ──
doc.add_heading("3.2 New Tables the Service Creates", level=2)

doc.add_paragraph(
    "The matching service creates and owns exactly 3 new tables in the same database. "
    "These tables store the AI-generated data: embeddings (numerical representations of "
    "text) and match scores (the results of the matching pipeline). They are managed by "
    "a separate migration chain so they never interfere with SMO's tables."
)

doc.add_heading("Table 1: candidate_embeddings", level=3)

doc.add_paragraph(
    "When the matching service processes a candidate, it takes their full resume text and "
    "sends it to OpenAI's embedding API. OpenAI returns a list of 1,536 numbers — a "
    "mathematical fingerprint that captures the meaning of the resume. Two similar resumes "
    "will have similar fingerprints, which is how the service finds good matches quickly "
    "without reading every job description from scratch."
)

doc.add_paragraph(
    "Each candidate gets exactly one row in this table. The embedding is updated only when "
    "the resume text changes (detected by comparing a hash of the text). This avoids "
    "unnecessary API calls — if a candidate's resume hasn't changed, their embedding is "
    "reused as-is."
)

add_table(
    ["Column", "Type", "Constraints", "What It Stores"],
    [
        ["id", "Integer", "PK, auto-increment", "Internal row ID (not meaningful outside the database)"],
        ["candidate_id", "UUID", "NOT NULL, UNIQUE, INDEX",
         "The candidate this embedding belongs to. References users.id from the SMO users table. "
         "UNIQUE means one embedding per candidate."],
        ["embedding_vector", "Vector(1536)", "NOT NULL",
         "The AI-generated fingerprint — 1,536 floating-point numbers produced by OpenAI's "
         "text-embedding-3-small model. Stored using the pgvector extension."],
        ["source_text_hash", "String(64)", "NOT NULL",
         "A SHA-256 hash of the resume text that was used to generate this embedding. On the "
         "next sync, if the hash matches, the embedding is skipped (no API call needed)."],
        ["model_name", "String(50)", "NOT NULL",
         "Which AI model produced this embedding (currently always 'text-embedding-3-small'). "
         "Tracked in case the model is upgraded in the future."],
        ["created_at", "Timestamp", "auto", "When this embedding was first generated."],
        ["updated_at", "Timestamp", "auto", "When this embedding was last refreshed (if the resume changed)."],
    ],
    col_widths=[2.5, 2, 3, 8],
)

add_note(
    "This table has an HNSW index on embedding_vector (cosine similarity, m=16, "
    "ef_construction=64). This is a specialized index for fast nearest-neighbor search — "
    "it lets the database find the 50 most similar jobs for a candidate in milliseconds, "
    "even with tens of thousands of embeddings."
)

doc.add_heading("Table 2: job_embeddings", level=3)

doc.add_paragraph(
    "The same concept as candidate_embeddings, but for job descriptions. Each job posting's "
    "description text is sent to OpenAI, and the resulting 1,536-number fingerprint is stored "
    "here. When matching, the service compares a candidate's embedding against all job "
    "embeddings to find the most similar ones — this is the 'semantic search' step of the "
    "pipeline."
)

add_table(
    ["Column", "Type", "Constraints", "What It Stores"],
    [
        ["id", "Integer", "PK, auto-increment", "Internal row ID"],
        ["job_id", "Integer", "NOT NULL, UNIQUE, INDEX",
         "The job this embedding belongs to. References jobs.id from the SMO jobs table. "
         "UNIQUE means one embedding per job."],
        ["embedding_vector", "Vector(1536)", "NOT NULL",
         "AI-generated fingerprint of the job description. Same model and dimensions as "
         "candidate embeddings, ensuring they can be compared directly."],
        ["source_text_hash", "String(64)", "NOT NULL",
         "Hash of the job description text. Skip re-embedding if the description hasn't changed."],
        ["model_name", "String(50)", "NOT NULL",
         "Which AI model was used (text-embedding-3-small)."],
        ["created_at", "Timestamp", "auto", "When first generated."],
        ["updated_at", "Timestamp", "auto", "When last refreshed."],
    ],
    col_widths=[2.5, 2, 3, 8],
)

add_note(
    "Same HNSW index as candidate_embeddings. Both tables use the same embedding model "
    "and dimensions (1,536), which is required — you can only compare vectors that were "
    "generated by the same model."
)

doc.add_heading("Table 3: match_scores_v2", level=3)

doc.add_paragraph(
    "This is where the final results live. Every time the matching pipeline scores a "
    "candidate-job pair, the result is stored here as a single row. The row includes "
    "the numerical scores (semantic similarity, LLM analysis, combined), plus a JSONB "
    "field called score_breakdown that contains the full 'WhyCard' — the AI's plain-English "
    "explanation of why this match is good or bad, what skills align, what gaps exist, and "
    "a recommendation."
)

doc.add_paragraph(
    "This table also doubles as the database cache layer. Each row has an expires_at "
    "timestamp set to 24 hours after creation. When the same match is requested again "
    "within that window, the cached result is returned instantly without re-running the "
    "pipeline. After expiration, the next request triggers a fresh analysis."
)

add_table(
    ["Column", "Type", "Constraints", "What It Stores"],
    [
        ["id", "Integer", "PK, auto-increment", "Internal row ID"],
        ["match_type", "String(20)", "NOT NULL",
         "Either 'candidate_to_jobs' (find jobs for a candidate) or 'job_to_candidates' "
         "(find candidates for a job). The same pair can have both types."],
        ["candidate_id", "UUID", "NOT NULL, INDEX",
         "Which candidate was matched. References users.id."],
        ["job_id", "Integer", "NOT NULL, INDEX",
         "Which job was matched. References jobs.id."],
        ["semantic_score", "Float", "NOT NULL",
         "How similar the resume and job description are, based on embedding comparison. "
         "Ranges from 0 (completely unrelated) to 1 (nearly identical meaning)."],
        ["rule_score", "Float", "NOT NULL",
         "Legacy field from a previous version of the pipeline. Always 0.0. Kept for "
         "schema compatibility — may be removed in a future migration."],
        ["llm_score", "Float", "Nullable",
         "The AI's assessment of fit after reading both the resume and job description. "
         "Ranges from 0 to 1. Null if this pair wasn't scored by the LLM (only top "
         "candidates/jobs go through LLM scoring)."],
        ["combined_score", "Float", "NOT NULL",
         "The final score shown to users: (semantic_score x 50%) + (llm_score x 50%). "
         "Weights are configurable. This is what results are ranked by."],
        ["score_breakdown", "JSONB", "NOT NULL",
         "The full WhyCard in JSON format. Contains: match_summary, strengths (top 3), "
         "concerns, skill_gaps, experience_fit, location_fit, salary_fit, and "
         "recommendation. This is rendered as the explanation card in the dashboard."],
        ["created_at", "Timestamp", "auto",
         "When this match was computed."],
        ["expires_at", "Timestamp", "INDEX",
         "Cache expiration: created_at + 24 hours. Rows past this time are ignored on "
         "read and eventually overwritten."],
    ],
    col_widths=[2.5, 1.8, 2.5, 8.7],
)

add_note(
    "Unique constraint on (match_type, candidate_id, job_id) — ensures there is only one "
    "active result per candidate-job pair per direction. When a fresh result is computed, "
    "it replaces (upserts) the existing row."
)


# ── 3.3 HOW NEW TABLES CONNECT ──
doc.add_heading("3.3 How New Tables Connect to Existing Tables", level=2)

doc.add_paragraph(
    "This is the most important part of this document. The 3 new tables don't exist in "
    "isolation — they reference data from existing SMO tables through their candidate_id "
    "and job_id columns. Here's how:"
)

doc.add_heading("candidate_embeddings", level=3)
doc.add_paragraph(
    "candidate_embeddings.candidate_id (UUID) points to users.id in the SMO users table. "
    "The matching service reads the candidate's profile from the users table and the "
    "candidate_profiles table (joined via candidate_profiles.user_id → users.id), generates "
    "an AI embedding from their resume and skills, and stores the result here. Each candidate "
    "gets exactly one row — when a new candidate appears in SMO, the matching service creates "
    "their embedding on the next sync."
)

doc.add_heading("job_embeddings", level=3)
doc.add_paragraph(
    "job_embeddings.job_id (Integer) points to jobs.id in the SMO jobs table. The matching "
    "service reads the job description from the jobs table (which links to employers via "
    "jobs.employer_id → employers.id), generates an AI embedding from the description text, "
    "and stores it here. Each job gets exactly one row."
)

doc.add_heading("match_scores_v2", level=3)
doc.add_paragraph(
    "match_scores_v2 has two foreign references:"
)
add_bullet(
    "references users.id — identifies which candidate was matched",
    bold_prefix="candidate_id (UUID):",
)
add_bullet(
    "references jobs.id — identifies which job was matched",
    bold_prefix="job_id (Integer):",
)
doc.add_paragraph(
    "This table is the result of matching — it connects a candidate to a job along with "
    "scores and the AI explanation (WhyCard). A single candidate can appear in many rows "
    "(one per job matched against), and a single job can appear in many rows (one per "
    "candidate matched against). For example, matching 100 candidates against 200 jobs "
    "could produce up to 20,000 rows."
)

doc.add_heading("Why No Hard Foreign Key Constraints?", level=3)
doc.add_paragraph(
    "You may notice that the new tables don't declare actual ForeignKey() constraints in "
    "the database. The values in candidate_id and job_id are correct (they match users.id "
    "and jobs.id), but the database does not enforce this at the schema level. This is a "
    "deliberate design choice, not an oversight. There are three reasons:"
)
add_bullet(
    "The SMO backend and the matching service each manage their database schema through "
    "separate Alembic migration chains. SMO tracks its version in the alembic_version "
    "table; the matching service tracks its version in alembic_version_auto_matcher. If "
    "the new tables declared hard FK constraints pointing at SMO tables, Alembic would "
    "try to manage tables it doesn't own, causing migration conflicts.",
    bold_prefix="Separate migration chains —",
)
add_bullet(
    "Because there are no hard FK constraints, either service can be deployed, rolled "
    "back, or restarted independently. You can drop and recreate all 3 matching tables "
    "without affecting any SMO data, and you can update SMO's schema without worrying "
    "about breaking the matching service's foreign key references.",
    bold_prefix="Decoupled deployments —",
)
add_bullet(
    "The matching service uses read-only SQLAlchemy mappings (defined in "
    "app/models/readonly.py with extend_existing=True) for all SMO tables. It physically "
    "cannot write to or modify SMO data, even accidentally. The relationship between the "
    "new tables and existing tables is enforced at the application layer (Python code), "
    "not at the database layer.",
    bold_prefix="Read-only contract —",
)

add_note(
    "In practice, this means if a candidate is deleted in SMO, their embedding row in "
    "candidate_embeddings becomes orphaned — but this is harmless. The matching service "
    "simply ignores orphaned rows. No errors, no cascading failures."
)

doc.add_heading("The Complete Reference Chain", level=3)
doc.add_paragraph(
    "Here is the full picture of how every table connects. Lines with a solid arrow (→) "
    "are hard FK constraints enforced by SMO's database schema. Lines with a dashed arrow "
    "(⇢) are logical references enforced by application code in the matching service:"
)

add_code_block(
    "EXISTING SMO TABLES (hard FK constraints)              NEW MATCHING TABLES (logical references)\n"
    "═══════════════════════════════════════════             ════════════════════════════════════════\n"
    "\n"
    "users.id ←── candidate_profiles.user_id (hard FK)      users.id ⇠── candidate_embeddings.candidate_id\n"
    "                                                        users.id ⇠── match_scores_v2.candidate_id\n"
    "\n"
    "jobs.id ──→ jobs.employer_id ──→ employers.id (hard FK) jobs.id  ⇠── job_embeddings.job_id\n"
    "                                                        jobs.id  ⇠── match_scores_v2.job_id\n"
    "\n"
    "users.pod_id ──→ pods.id (hard FK)\n"
    "candidate_profiles.university_id ──→ universities.id (hard FK)"
)

doc.add_paragraph(
    "When the dashboard displays a match result, it follows these chains to assemble "
    "the full picture. For example, to show 'John Doe matched to Software Engineer at "
    "Acme Corp, score 85/100', the system:"
)

add_bullet(
    "Reads the match score row from match_scores_v2 (gets the score and WhyCard)",
    bold_prefix="Step 1:",
)
add_bullet(
    "Follows candidate_id → users.id to get the candidate's name and email",
    bold_prefix="Step 2:",
)
add_bullet(
    "Follows users.id → candidate_profiles.user_id to get resume, skills, experience",
    bold_prefix="Step 3:",
)
add_bullet(
    "Follows job_id → jobs.id to get the job title and description",
    bold_prefix="Step 4:",
)
add_bullet(
    "Follows jobs.employer_id → employers.id to get the company name",
    bold_prefix="Step 5:",
)

add_diagram(diagrams["db_ownership"])

doc.add_heading("Foreign Key Summary Table", level=3)

doc.add_paragraph(
    "For quick reference, here is every relationship between the new and existing tables:"
)

add_table(
    ["New Table", "Column", "Points To", "Type", "What It Means"],
    [
        ["candidate_embeddings", "candidate_id\n(UUID)", "users.id",
         "One-to-one\n(logical)",
         "Each candidate gets exactly one embedding row. The matching "
         "service looks up this user in SMO to get their resume text for "
         "AI processing."],
        ["job_embeddings", "job_id\n(Integer)", "jobs.id",
         "One-to-one\n(logical)",
         "Each job gets exactly one embedding row. The matching service "
         "looks up this job in SMO to get the description text for AI processing."],
        ["match_scores_v2", "candidate_id\n(UUID)", "users.id",
         "Many-to-one\n(logical)",
         "One candidate can have many match score rows — one for each "
         "job they were matched against. This is the candidate side of "
         "every match result."],
        ["match_scores_v2", "job_id\n(Integer)", "jobs.id",
         "Many-to-one\n(logical)",
         "One job can have many match score rows — one for each candidate "
         "it was matched against. This is the job side of every match result."],
    ],
    col_widths=[2.5, 2, 1.5, 2, 7.5],
)


# ── 3.4 COMPLETE DATA CHAIN ──
doc.add_heading("3.4 Complete Data Chain (Worked Example)", level=2)

doc.add_paragraph(
    "To make the relationships concrete, here's a step-by-step walkthrough of what happens "
    "when the system matches a real candidate to a real job — showing which tables are read, "
    "which are written, and what data flows between them."
)

doc.add_paragraph(
    "In this example, candidate 'John Doe' (a Java developer with 3 years of experience, "
    "based in NYC) is being matched against a 'Software Engineer' position at Acme Corp "
    "(located in NYC, $90-120K salary, 2 years experience required)."
)

add_diagram(diagrams["match_chain"])

doc.add_paragraph(
    "At the end of this process, the consultant sees a match score of 80/100 on the "
    "dashboard with a WhyCard that explains: strong Java alignment, Python is a bonus, "
    "but there's a gap in AWS experience. The consultant can use this information to "
    "decide whether to submit John for this position, and if so, to coach John on "
    "highlighting any cloud experience he may have."
)


# ── 3.5 DB OWNERSHIP & MIGRATION SAFETY ──
doc.add_heading("3.5 Database Ownership & Migration Safety", level=2)

doc.add_paragraph(
    "A common concern when two services share a database: how do you make sure they don't "
    "break each other's tables? The answer is separate migration chains with strict "
    "ownership boundaries."
)

doc.add_paragraph(
    "Both services use Alembic (Python's standard database migration tool) but they each "
    "maintain their own version tracking table. This means you can run migrations for either "
    "service at any time without affecting the other. The matching service includes an "
    "additional safety filter (OWNED_TABLES) that explicitly restricts its migration "
    "autogeneration to only its 3 tables — even if it can see all 36 tables in the database, "
    "it will never generate migration scripts that touch SMO's tables."
)

add_table(
    ["Aspect", "SMO Backend", "Matching Service"],
    [
        ["Migration tool", "Alembic", "Alembic"],
        ["Version table", "alembic_version", "alembic_version_auto_matcher"],
        ["Tables managed", "33 tables (all SMO data)",
         "3 tables only (candidate_embeddings, job_embeddings, match_scores_v2)"],
        ["Safety filter", "None needed (owns the whole database)",
         "OWNED_TABLES filter in env.py — autogenerate ignores all other tables"],
        ["Run independently?", "Yes",
         "Yes — migrations never touch each other's tables"],
    ],
    col_widths=[3, 5.5, 7],
)

add_code_block(
    "# Matching service migration commands\n"
    "cd auto-matching-service\n"
    "alembic upgrade head              # Apply migrations (only creates/modifies 3 owned tables)\n"
    "alembic revision --autogenerate   # Generate new migration (only detects changes in 3 tables)"
)

add_note(
    "You can safely run 'alembic upgrade head' for both services in any order. They use "
    "separate version tables, so they track their migration history independently. Running "
    "one will never roll back or interfere with the other."
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════
# 4. AI MATCHING PIPELINE
# ══════════════════════════════════════════════════════════
doc.add_heading("4. AI Matching Pipeline", level=1)

doc.add_paragraph(
    "The matching pipeline is the core logic that takes a candidate and finds their best "
    "job matches (or vice versa). It works as a funnel: start with all available jobs, "
    "progressively filter and rank them, and produce a short list of the best matches "
    "with detailed explanations."
)


# ── 4.1 THREE-LAYER PIPELINE ──
doc.add_heading("4.1 Three-Layer Pipeline", level=2)

doc.add_paragraph(
    "The pipeline has three layers, each serving a different purpose. The first layer is "
    "cheap and fast (pure SQL), the second is medium-cost (vector math), and the third "
    "is expensive but intelligent (AI analysis). By running them in order, the system "
    "avoids wasting expensive AI calls on obviously bad matches."
)

add_table(
    ["Layer", "What It Does", "Data Used", "Output"],
    [
        ["1. Hard Filter\n(SQL, ~100ms)",
         "Eliminates obviously bad matches using simple rules. A job requiring "
         "5 years of experience won't be shown to a candidate with 2 years. "
         "A blacklisted employer is excluded entirely. This is fast and free — "
         "it's just a database query.",
         "candidate_profiles\n(experience, location, salary)\n\n"
         "jobs\n(experience, location, salary)\n\n"
         "employers\n(do_not_apply flag)",
         "Up to 8,000 viable jobs\n(from tens of thousands)"],
        ["2. Semantic Search\n(pgvector, ~200ms)",
         "Compares the meaning of the candidate's resume against every surviving "
         "job description. Uses the pre-computed embeddings and a fast nearest-neighbor "
         "index. This finds non-obvious matches that keyword search would miss — e.g., "
         "a 'Python developer' resume matching a job that asks for 'Django backend engineer'.",
         "candidate_embeddings\n.embedding_vector\n\n"
         "job_embeddings\n.embedding_vector",
         "Top 50 most similar jobs,\neach with a semantic score\n(0.0 to 1.0)"],
        ["3. LLM Analysis\n(Cerebras, ~2-3s each)",
         "For the top candidates/jobs, sends both sides to an AI model that reads "
         "the resume and job description in detail. The AI produces a structured "
         "analysis: overall score, matching skills, skill gaps, experience fit, "
         "salary fit, concerns, and a recommendation. This is the 'WhyCard' that "
         "consultants see.",
         "candidate name, skills,\nexperience, location, salary\n\n"
         "job title, description,\nlocation, salary, experience\n\n"
         "employer name",
         "Top 10 scored matches,\neach with a WhyCard:\nscore, strengths,\nskill gaps, "
         "recommendation"],
    ],
    col_widths=[2.5, 5, 4, 4],
)

doc.add_paragraph(
    "The final score combines the semantic and LLM scores using configurable weights:"
)

add_code_block(
    "combined_score = (WEIGHT_SEMANTIC x semantic_score) + (WEIGHT_LLM x llm_score)\n"
    "                  default: 0.5                        default: 0.5"
)

doc.add_paragraph(
    "These weights can be adjusted from the dashboard settings page in real-time — "
    "no code changes or restarts needed. For example, if you trust the AI's judgment "
    "more than the raw text similarity, you could set WEIGHT_LLM to 0.7 and "
    "WEIGHT_SEMANTIC to 0.3."
)


# ── 4.2 TWO-LAYER CACHING ──
doc.add_heading("4.2 Two-Layer Caching", level=2)

doc.add_paragraph(
    "The LLM analysis step (Layer 3) is the slowest and most expensive part of the "
    "pipeline — each Cerebras API call takes 2-3 seconds and has a per-call cost. "
    "Matching 100 candidates against 200 jobs means up to 20,000 LLM calls. Without "
    "caching, repeating the same match request would redo all that work."
)

doc.add_paragraph(
    "To solve this, the service uses two layers of caching. The first layer (database) "
    "caches complete match results — if you request the same match again within 24 hours, "
    "the entire pipeline is skipped and the cached result is returned in ~5ms. The second "
    "layer (Redis) caches individual LLM results — if a specific candidate-job pair was "
    "already analyzed, the Cerebras API call is skipped even if the overall request is "
    "different."
)

add_diagram(diagrams["cache_flow"])

add_table(
    ["Cache Layer", "Where", "Cache Key", "TTL", "What's Cached", "When It Helps"],
    [
        ["Layer 1:\nDB Cache",
         "match_scores_v2\ntable",
         "(match_type,\ncandidate_id,\njob_id)",
         "24\nhours",
         "Full match result\nwith WhyCard",
         "Exact same match request repeated.\n"
         "Skips entire pipeline. ~5ms response."],
        ["Layer 2:\nRedis",
         "Redis\n(in-memory)",
         "llm_score:\n{candidate_id}:\n{job_id}",
         "72\nhours",
         "Individual LLM\nresult (score +\nanalysis)",
         "Same pair appears in different request.\n"
         "E.g., 'jobs for John' cached pair reused\n"
         "in 'candidates for Acme SWE'. ~1ms."],
    ],
    col_widths=[1.8, 2, 2.5, 1, 3, 5.2],
)

add_note(
    "Redis is fully optional. If REDIS_URL is not configured or Redis goes down at runtime, "
    "the service logs a warning and continues with database-only caching. No code changes, "
    "no restarts, no errors — it self-heals when Redis comes back."
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════
# 5. DEPLOYMENT
# ══════════════════════════════════════════════════════════
doc.add_heading("5. Deployment", level=1)


# ── 5.1 ENV VARS ──
doc.add_heading("5.1 Environment Variables", level=2)

doc.add_paragraph(
    "The service is configured entirely through environment variables. Only 3 are required — "
    "everything else has sensible defaults. All variables are loaded through Pydantic Settings "
    "in app/core/config.py."
)

doc.add_heading("Required", level=3)

add_table(
    ["Variable", "Example", "Description"],
    [
        ["DATABASE_URL", "postgresql+asyncpg://user@host:5432/smo_local",
         "Connection string for the shared PostgreSQL database. Must use the asyncpg "
         "driver. This should point to the same database the SMO backend uses."],
        ["OPENAI_API_KEY", "sk-...",
         "API key for OpenAI. Used to generate text embeddings (text-embedding-3-small). "
         "Embedding calls are inexpensive (~$0.0001 per document)."],
        ["CEREBRAS_API_KEY", "csk-...",
         "API key for Cerebras. Used for LLM scoring and analysis (gpt-oss-120b). "
         "This is the primary API cost — ~$0.0001-0.0005 per call."],
    ],
    col_widths=[3, 5, 7.5],
)

doc.add_heading("Optional (with defaults)", level=3)

add_table(
    ["Variable", "Default", "What It Controls"],
    [
        ["REDIS_URL", "(empty — disabled)",
         "Redis connection string for LLM caching. Leave empty to run without Redis."],
        ["PORT", "8002", "Backend server port."],
        ["FRONTEND_URL", "http://localhost:3001", "Dashboard URL, used for CORS."],
        ["WEIGHT_SEMANTIC", "0.5", "How much the semantic score contributes to the final score."],
        ["WEIGHT_LLM", "0.5", "How much the LLM score contributes. Must sum to 1.0 with above."],
        ["MATCH_CACHE_TTL_HOURS", "24", "How long full match results are cached in the database."],
        ["LLM_CACHE_TTL_HOURS", "72", "How long individual LLM results are cached in Redis."],
        ["LLM_CONCURRENCY_LIMIT", "5", "Max parallel Cerebras API calls. Higher = faster but may hit rate limits."],
        ["JOB_RECENCY_DAYS", "30", "Only consider jobs posted within this many days."],
        ["EMBEDDING_BATCH_SIZE", "64", "How many documents to embed per OpenAI API call."],
    ],
    col_widths=[3.5, 3, 9],
)


# ── 5.2 LOCAL DEV ──
doc.add_heading("5.2 Local Development", level=2)

doc.add_paragraph(
    "To run the matching service locally, you need the existing PostgreSQL database "
    "(smo_local) running, plus API keys for OpenAI and Cerebras. Redis is optional."
)

doc.add_heading("Step 1: Start the backend", level=3)
add_code_block(
    "cd auto-matching-service\n"
    "cp .env.example .env          # Edit .env with your API keys\n"
    "pip install -r requirements.txt\n"
    "alembic upgrade head           # Create the 3 new tables\n"
    "uvicorn app.main:app --reload --port 8002"
)

add_note(
    "The 'alembic upgrade head' command creates candidate_embeddings, job_embeddings, and "
    "match_scores_v2 tables. It will not modify any existing SMO tables."
)

doc.add_heading("Step 2: Start the dashboard", level=3)
add_code_block(
    "cd auto-matching-service/frontend\n"
    "npm install\n"
    "npm run dev                    # Starts on port 3001"
)

doc.add_heading("Step 3 (optional): Start Redis", level=3)
add_code_block(
    "# Add to your .env file:\n"
    "REDIS_URL=redis://localhost:6379/0\n\n"
    "# Then start Redis:\n"
    "redis-server"
)

doc.add_heading("Verification", level=3)
doc.add_paragraph(
    "Once running, verify the service is healthy:"
)
add_bullet("Backend health check: GET http://localhost:8002/api/v1/health")
add_bullet("Dashboard: Open http://localhost:3001 in your browser")
add_bullet("Login: admin / password@123")


# ── 5.3 DOCKER & CLOUD RUN ──
doc.add_heading("5.3 Docker & Cloud Run", level=2)

doc.add_paragraph(
    "For production deployment, the service is containerized with Docker and deployed to "
    "Google Cloud Run. Cloud Run scales to zero when idle (no traffic = no cost) and "
    "auto-scales under load."
)

doc.add_heading("Backend Dockerfile", level=3)
add_code_block(
    "FROM python:3.11-slim\n"
    "WORKDIR /app\n"
    "COPY requirements.txt .\n"
    "RUN pip install --no-cache-dir -r requirements.txt\n"
    "COPY . .\n"
    'CMD ["uvicorn", "app.main:app", "--host", "0.0.0.0", "--port", "8080"]'
)

doc.add_heading("Frontend Dockerfile", level=3)
add_code_block(
    "FROM node:20-slim AS builder\n"
    "WORKDIR /app\n"
    "COPY package.json package-lock.json ./\n"
    "RUN npm ci\n"
    "COPY . .\n"
    "RUN npm run build\n\n"
    "FROM node:20-slim\n"
    "WORKDIR /app\n"
    "COPY --from=builder /app/.next/standalone ./\n"
    "COPY --from=builder /app/.next/static ./.next/static\n"
    "COPY --from=builder /app/public ./public\n"
    'CMD ["node", "server.js"]'
)

doc.add_heading("Deploy Commands", level=3)
add_code_block(
    "# Deploy backend to Cloud Run\n"
    "gcloud run deploy auto-matcher-backend \\\n"
    "  --source=auto-matching-service/ \\\n"
    "  --set-env-vars DATABASE_URL=...,OPENAI_API_KEY=...,CEREBRAS_API_KEY=... \\\n"
    "  --allow-unauthenticated --region us-central1 --memory 1Gi\n\n"
    "# Deploy frontend to Cloud Run\n"
    "gcloud run deploy auto-matcher-frontend \\\n"
    "  --source=auto-matching-service/frontend/ \\\n"
    "  --set-env-vars NEXT_PUBLIC_API_URL=https://auto-matcher-backend-xxx.run.app \\\n"
    "  --allow-unauthenticated --region us-central1"
)

doc.add_heading("Required GCP Services", level=3)
add_bullet(
    "The matching service needs to reach the same PostgreSQL instance that SMO uses. "
    "If SMO already runs on Cloud SQL, the matching service connects to it — no new "
    "database instance needed.",
    bold_prefix="Cloud SQL (PostgreSQL 15 + pgvector):"
)
add_bullet(
    "Optional but recommended for production. Provides fast LLM result caching. "
    "Basic M1 tier (1 GB) is sufficient.",
    bold_prefix="Memorystore Redis:"
)
add_bullet(
    "Two services: auto-matcher-backend and auto-matcher-frontend. Both scale to "
    "zero when idle.",
    bold_prefix="Cloud Run:"
)
add_bullet(
    "Store DATABASE_URL, OPENAI_API_KEY, and CEREBRAS_API_KEY as secrets rather "
    "than plain environment variables.",
    bold_prefix="Secret Manager:"
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════
# 6. COST ESTIMATES
# ══════════════════════════════════════════════════════════
doc.add_heading("6. Cost Estimates", level=1)

doc.add_paragraph(
    "The matching service has two cost components: infrastructure (servers, database, "
    "cache) and API usage (OpenAI and Cerebras calls). Infrastructure costs are fixed "
    "monthly regardless of usage. API costs depend on how many candidates and jobs are "
    "matched, but caching significantly reduces repeat costs."
)

doc.add_heading("Infrastructure (Monthly)", level=2)

add_table(
    ["Service", "Tier", "Estimated Cost"],
    [
        ["Cloud Run (backend)", "0.5 vCPU, 1 GB RAM, scales to zero when idle", "$5–15/mo"],
        ["Cloud Run (frontend)", "0.5 vCPU, 512 MB RAM, scales to zero when idle", "$3–8/mo"],
        ["Cloud SQL", "Shared with existing SMO instance — no new instance needed", "$0 incremental"],
        ["Memorystore Redis", "Basic M1, 1 GB (optional)", "$7–10/mo"],
        ["", "", ""],
        ["Total infrastructure", "", "$15–33/mo"],
    ],
    col_widths=[4, 6.5, 3],
)

doc.add_heading("API Costs (Per Batch Run)", level=2)

add_table(
    ["Scenario", "OpenAI Embeddings", "Cerebras LLM Calls", "Total API Cost"],
    [
        ["10 candidates x 20 jobs each",
         "~30 embeddings ($0.01)",
         "~200 LLM calls ($0.20)",
         "~$0.21"],
        ["100 candidates x 200 jobs each",
         "~300 embeddings ($0.05)",
         "~20,000 LLM calls ($1-5)",
         "~$1-5"],
        ["Repeat same batch (cached)",
         "$0 (embeddings reused)",
         "$0 (results cached in DB/Redis)",
         "$0"],
    ],
    col_widths=[4, 3.5, 4, 3],
)

doc.add_paragraph(
    "The caching system is key to cost control. After the first run for any set of "
    "candidate-job pairs, subsequent requests for the same pairs are served entirely "
    "from cache — no API calls, no cost. Database cache lasts 24 hours; Redis cache "
    "lasts 72 hours. For ongoing operations where the same candidates are matched "
    "against new jobs periodically, only the new pairs incur API costs."
)


# ══════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════
output_path = "Wynisco_Infrastructure_Document.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
